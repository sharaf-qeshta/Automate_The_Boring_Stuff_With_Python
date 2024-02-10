"""
Custom Invitations as Word Documents
Say you have a text file of guest names. This guests.txt file has one name per
line, as follows:
Prof. Plum
Miss Scarlet
Col. Mustard
Al Sweigart
RoboCop
Write a program that would generate a Word document with custom
invitations that look like Figure 15-11.
Since Python-Docx can use only those styles that already exist in the
Word document, you will have to first add these styles to a blank Word file
and then open that file with Python-Docx. There should be one invitation
per page in the resulting Word document, so call add_break() to add a page
break after the last paragraph of each invitation. This way, you will need to
open only one Word document to print all of the invitations at once.
You can download a sample guests.txt file from https://nostarch.com
/automatestuff2/

@author Sharaf Qeshta
"""

import docx

document = docx.Document("InvitationCard.docx")
names_file = open("guests.txt")

for line in names_file:
    document.add_paragraph('It would be a pleasure to have the company of').style = 'InvitationLine1'
    document.add_paragraph(line).style = 'InvitationName'
    document.add_paragraph('at 11010 Memory Lane on the Evening of').style = 'InvitationAddress'
    document.add_paragraph('April 1st').style = 'InvitationDate'
    document.add_paragraph('at 7 o\'clock').style = 'InvitationTime'
    document.paragraphs[4].runs[4].add_break(docx.enum.text.WD_BREAK.PAGE)

document.save("InvitationCard.docx")
